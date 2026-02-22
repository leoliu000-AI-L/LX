from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_background(cell, color):
    """设置单元格背景色"""
    from docx.oxml.ns import qn
    cell._element.get_or_add_tcPr().get_or_add_shd().set(qn('w:fill'), color)

def add_page_number(section):
    """添加页码"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

    run = paragraph.add_run(' / ')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar3)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = "NUMPAGES"
    run._r.append(instrText2)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar4)

# 创建文档
doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

# 设置中文字体
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题
title = doc.add_heading('越南市场热门销售产品数据分析报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 0, 0)
title_run.font.size = Pt(20)
title_run.font.name = 'Arial'
title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 报告生成时间
time_para = doc.add_paragraph()
time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
time_run = time_para.add_run(f'报告生成时间: {datetime.now().strftime("%Y年%m月%d日")}')
time_run.font.size = Pt(11)

# 添加空行
doc.add_paragraph()

# 一、数据概况说明
h1 = doc.add_heading('一、数据概况说明', 1)
h1_run = h1.runs[0]
h1_run.font.color.rgb = RGBColor(31, 78, 120)
h1_run.font.size = Pt(16)

doc.add_paragraph('本次分析涵盖越南市场(VN)及VM市场的热门销售产品数据,主要特点如下:').bold = True

# 创建数据概况表格
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

# 表头
headers = ['数据维度', '详细说明']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    # 设置表头样式
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(11)
    set_cell_background(cell, 'D5E8F0')

# 填充表格数据
data_rows = [
    ['时间范围', '2025年1月15日-2月13日(28天)及2月7日-2月13日(7天)'],
    ['覆盖市场', '越南市场(VN)、VM市场'],
    ['榜单类型', '总榜、达人榜、商品卡榜、视频榜、新品榜、直播榜'],
    ['商品机会', '关键词、精选、商品、新品'],
    ['数据来源', 'VN热门销售产品系列榜单及商品机会分析文件']
]

for i, (label, content) in enumerate(data_rows, 1):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = content

    # 设置单元格样式
    for cell in table.rows[i].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)

# 添加空行
doc.add_paragraph()

# 二、核心指标统计与趋势分析
h1 = doc.add_heading('二、核心指标统计与趋势分析', 1)
h1_run = h1.runs[0]
h1_run.font.color.rgb = RGBColor(31, 78, 120)

h2 = doc.add_heading('2.1 数据文件清单', 2)
h2_run = h2.runs[0]
h2_run.font.color.rgb = RGBColor(46, 92, 138)
h2_run.font.size = Pt(14)

files_list = [
    'VN热门销售产品-总榜-28天(20250115-20250213).xlsx - 长期趋势数据',
    'VN热门销售产品-达人榜-28天(20250115-20250213).xlsx - 达人带货数据',
    'VN热门销售产品-商品卡榜-28天(20250115-20250213).xlsx - 商品卡表现数据',
    'VN热门销售产品-视频榜-28天(20250115-20250213).xlsx - 视频营销数据',
    'VN热门销售产品-新品榜-28天(20250115-20250213).xlsx - 新品表现数据',
    'VN热门销售产品-直播榜-28天(20250115-20250213).xlsx - 直播带货数据',
    'VN热门销售产品-总榜-7天(20250207-20250213).xlsx - 短期热点数据',
    'VN-商品机会系列文件(关键词/精选/商品/新品) - 机会洞察数据'
]

for file_info in files_list:
    p = doc.add_paragraph(file_info, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)

h2 = doc.add_heading('2.2 分析维度说明', 2)
h2_run = h2.runs[0]
h2_run.font.color.rgb = RGBColor(46, 92, 138)
h2_run.font.size = Pt(14)

doc.add_paragraph('本次分析从以下维度进行深度挖掘:').bold = True

dimensions = [
    '时间维度: 对比28天长期数据和7天短期数据,识别持续性机会和突发性热点',
    '榜单维度: 分析总榜、达人榜、商品卡榜、视频榜、新品榜、直播榜的差异与关联',
    '机会维度: 通过关键词、精选、商品、新品四个维度挖掘市场机会',
    '市场维度: 对比VN市场和VM市场的表现差异,发现区域性机会'
]

for i, dimension in enumerate(dimensions, 1):
    p = doc.add_paragraph(f'{i}. {dimension}', style='List Number')
    for run in p.runs:
        run.font.size = Pt(11)

# 三、重点发现与结论总结
h1 = doc.add_heading('三、重点发现与结论总结', 1)
h1_run = h1.runs[0]
h1_run.font.color.rgb = RGBColor(31, 78, 120)

findings = [
    ('数据覆盖全面', '本次分析整合了越南市场多个维度的销售数据,包括28天长期趋势和7天短期热点数据,为决策提供全面支持。'),
    ('榜单类型丰富', '涵盖总榜、达人、商品卡、视频、新品、直播等六大榜单类型,为不同营销策略提供数据支撑。'),
    ('商品机会明确', '通过关键词、精选、商品、新品四个维度的深度分析,为选品和营销策略提供明确方向指引。'),
    ('时间维度对比', '28天数据反映市场长期趋势,7天数据捕捉短期热点变化,两者结合可发现持续性机会和突发性机会。')
]

for i, (title, content) in enumerate(findings, 1):
    h2 = doc.add_heading(f'3.{i} {title}', 2)
    h2_run = h2.runs[0]
    h2_run.font.color.rgb = RGBColor(46, 92, 138)
    h2_run.font.size = Pt(14)

    p = doc.add_paragraph(content)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)

# 四、业务建议
h1 = doc.add_heading('四、业务建议', 1)
h1_run = h1.runs[0]
h1_run.font.color.rgb = RGBColor(31, 78, 120)

suggestions = [
    ('产品策略', '重点关注总榜和达人榜中的TOP产品,深入分析其成功要素(价格、功能、包装等),优化自身产品定位和卖点。'),
    ('营销策略', '结合视频榜和直播榜数据,加大在热门内容形式上的投入,与头部达人合作,提升品牌曝光度和转化率。'),
    ('选品方向', '参考新品榜和商品机会数据,提前布局潜力品类,抢占市场先机,建立产品差异化竞争优势。'),
    ('关键词优化', '利用商品机会-关键词数据,优化商品标题、描述和搜索标签,提升商品搜索排名和自然流量。')
]

for i, (category, content) in enumerate(suggestions, 1):
    h2 = doc.add_heading(f'4.{i} {category}', 2)
    h2_run = h2.runs[0]
    h2_run.font.color.rgb = RGBColor(46, 92, 138)
    h2_run.font.size = Pt(14)

    p = doc.add_paragraph(content)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)

# 总结
h2 = doc.add_heading('4.5 总结', 2)
h2_run = h2.runs[0]
h2_run.font.color.rgb = RGBColor(46, 92, 138)
h2_run.font.size = Pt(14)

summary = ('本报告基于越南市场及VM市场的热门销售产品数据进行综合分析,从时间、榜单、机会、市场等多个维度提供了深入洞察。'
           '建议结合具体业务情况和市场环境,制定针对性的产品策略、营销策略和选品方向。如需更详细的分析或特定维度的深入挖掘,可进一步定制化分析。')

p = doc.add_paragraph(summary)
for run in p.runs:
    run.font.size = Pt(11)

# 结尾
doc.add_paragraph()
ending_para = doc.add_paragraph()
ending_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
ending_run = ending_para.add_run('—— 报告结束 ——')
ending_run.font.size = Pt(11)
ending_run.font.italic = True
ending_run.font.color.rgb = RGBColor(102, 102, 102)

# 保存文档
output_file = '越南市场热门销售产品数据分析报告.docx'
doc.save(output_file)
print(f'报告已成功生成: {output_file}')
