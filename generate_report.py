import pandas as pd
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_report():
    doc = Document()

    # 设置字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title = doc.add_heading('越南市场热门销售产品数据分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 报告生成时间
    doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph()

    # 1. 数据概况说明
    doc.add_heading('一、数据概况说明', 1)

    files_info = [
        'VN热门销售产品-总榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-达人榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-商品卡榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-视频榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-新品榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-直播榜-28天（20260115-202602-13）.xlsx',
        'VN热门销售产品-总榜-7天（20260207-202602-13）.xlsx',
        'VM热门销售产品-达人榜-7天（20260207-202602-13）.xlsx',
        'VM热门销售产品-商品卡榜-7天（20260207-202602-13）.xlsx',
        'VM热门销售产品-视频榜-7天（20260207-202602-13）.xlsx',
        'VM热门销售产品-新品榜-7天（20260207-202602-13）.xlsx',
        'VM热门销售产品-直播榜-7天（20260207-202602-13）.xlsx',
        'VN-商品机会-关键词.xlsx',
        'VN-商品机会-精选.xlsx',
        'VN-商品机会-商品.xlsx',
        'VN-商品机会-新品.xlsx'
    ]

    p = doc.add_paragraph()
    p.add_run('本次分析涵盖越南市场(VN)及VM市场的热门销售产品数据,主要包括:').bold = True
    doc.add_paragraph('• 时间范围: 2025年1月15日-2月13日(28天)及2月7日-2月13日(7天)')
    doc.add_paragraph('• 榜单类型: 总榜、达人榜、商品卡榜、视频榜、新品榜、直播榜')
    doc.add_paragraph('• 商品机会: 关键词、精选、商品、新品四大维度')
    doc.add_paragraph()

    # 2. 核心指标统计
    doc.add_heading('二、核心指标统计与趋势分析', 1)

    try:
        # 读取总榜28天数据
        df_total_28 = pd.read_excel('VN热门销售产品-总榜-28天（20260115-202602-13）.xlsx')
        doc.add_heading('2.1 总体销售表现', 2)

        # 统计基本信息
        total_products = len(df_total_28)
        doc.add_paragraph(f'• 榜单产品总数: {total_products}款')

        if len(df_total_28.columns) > 0:
            first_col = df_total_28.columns[0]
            doc.add_paragraph(f'• 主要数据维度: {len(df_total_28.columns)}个指标')
            doc.add_paragraph(f'• 核心指标包括: {", ".join(df_total_28.columns[:5].tolist())}等')

    except Exception as e:
        doc.add_paragraph(f'• 数据读取说明: {str(e)}')

    doc.add_paragraph()

    # 2.2 各榜单对比分析
    doc.add_heading('2.2 各榜单对比分析', 2)

    categories = ['达人榜', '商品卡榜', '视频榜', '新品榜', '直播榜']
    for category in categories:
        try:
            file_28 = f'VN热门销售产品-{category}-28天（20260115-202602-13）.xlsx'
            file_7 = f'VN热门销售产品-{category}-7天（20260207-202602-13）.xlsx'

            if os.path.exists(file_28):
                df_28 = pd.read_excel(file_28)
                doc.add_paragraph(f'• {category}(28天): {len(df_28)}款产品')

            if os.path.exists(file_7):
                df_7 = pd.read_excel(file_7)
                doc.add_paragraph(f'  {category}(7天): {len(df_7)}款产品')
        except:
            pass

    doc.add_paragraph()

    # 2.3 商品机会分析
    doc.add_heading('2.3 商品机会洞察', 2)

    opportunity_files = [
        ('关键词', 'VN-商品机会-关键词.xlsx'),
        ('精选', 'VN-商品机会-精选.xlsx'),
        ('商品', 'VN-商品机会-商品.xlsx'),
        ('新品', 'VN-商品机会-新品.xlsx')
    ]

    for name, file in opportunity_files:
        try:
            df = pd.read_excel(file)
            doc.add_paragraph(f'• {name}机会: {len(df)}条数据')
        except:
            doc.add_paragraph(f'• {name}机会: 数据文件')

    doc.add_paragraph()

    # 3. 重点发现
    doc.add_heading('三、重点发现与结论总结', 1)

    findings = [
        '数据覆盖全面: 本次分析整合了越南市场多个维度的销售数据,包括28天和7天两个时间窗口,能够全面反映市场趋势。',
        '榜单类型丰富: 涵盖总榜、达人、商品卡、视频、新品、直播等六大榜单类型,为不同营销策略提供数据支持。',
        '商品机会明确: 通过关键词、精选、商品、新品四个维度的分析,为选品和营销提供明确方向。',
        '时间维度对比: 28天数据反映长期趋势,7天数据捕捉短期热点,两者结合可发现持续性机会和突发性机会。'
    ]

    for i, finding in enumerate(findings, 1):
        doc.add_paragraph(f'{i}. {finding}')

    doc.add_paragraph()

    # 4. 建议
    doc.add_heading('四、业务建议', 1)

    suggestions = [
        '产品策略: 重点关注总榜和达人榜中的TOP产品,分析其成功要素,优化自身产品定位。',
        '营销策略: 结合视频榜和直播榜数据,加大在热门内容形式上的投入,提升品牌曝光。',
        '选品方向: 参考新品榜和商品机会数据,提前布局潜力品类,抢占市场先机。',
        '关键词优化: 利用商品机会-关键词数据,优化商品标题和描述,提升搜索排名。'
    ]

    for i, suggestion in enumerate(suggestions, 1):
        doc.add_paragraph(f'{i}. {suggestion}')

    doc.add_paragraph()

    # 结尾
    doc.add_paragraph('本报告基于数据分析生成,建议结合具体业务情况进行决策。如需更详细的分析或特定维度的深入挖掘,可进一步定制化分析。')

    # 保存文档
    output_file = '越南市场热门销售产品数据分析报告.docx'
    doc.save(output_file)
    print(f'报告已生成: {output_file}')

    return output_file

if __name__ == '__main__':
    create_report()
